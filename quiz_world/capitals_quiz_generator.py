import random
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt
from reportlab.pdfgen import canvas


# World capitals dataset (country -> capital)
capitals = {
    'afghanistan': 'kabul', 'albania': 'tirana', 'germany': 'berlin',
    'andorra': 'andorra la vella', 'angola': 'luanda', 'antigua and barbuda': "saint john's",
    'saudi arabia': 'riyadh', 'algeria': 'algiers', 'argentina': 'buenos aires',
    'armenia': 'yerevan', 'australia': 'canberra', 'austria': 'vienna',
    'azerbaijan': 'baku', 'bahamas': 'nassau', 'barbados': 'bridgetown',
    'bahrain': 'manama', 'belgium': 'brussels', 'belize': 'belmopan',
    'benin': 'porto-novo', 'belarus': 'minsk', 'botswana': 'gaborone',
    'brazil': 'brasilia', 'brunei': 'bandar seri begawan', 'bulgaria': 'sofia',
    'burkina faso': 'ouagadougou', 'burundi': 'gitega', 'bhutan': 'thimphu',
    'cabo verde': 'praia', 'cambodia': 'phnom penh', 'cameroon': 'yaounde',
    'canada': 'ottawa', 'qatar': 'doha', 'chad': "n'djamena",
    'chile': 'santiago', 'china': 'beijing', 'cyprus': 'nicosia',
    'colombia': 'bogota', 'comoros': 'moroni', 'republic of the congo': 'brazzaville',
    'democratic republic of the congo': 'kinshasa', 'north korea': 'pyongyang',
    'south korea': 'seoul', "cote d'ivoire": 'yamoussoukro', 'costa rica': 'san jose',
    'croatia': 'zagreb', 'cuba': 'havana', 'denmark': 'copenhagen',
    'djibouti': 'djibouti', 'dominica': 'roseau', 'dominican republic': 'santo domingo',
    'ecuador': 'quito', 'egypt': 'cairo', 'el salvador': 'san salvador',
    'united arab emirates': 'abu dhabi', 'eritrea': 'asmara', 'slovakia': 'bratislava',
    'slovenia': 'ljubljana', 'spain': 'madrid', 'estonia': 'tallinn',
    'ethiopia': 'addis ababa', 'philippines': 'manila', 'finland': 'helsinki',
    'france': 'paris', 'gabon': 'libreville', 'gambia': 'banjul',
    'georgia': 'tbilisi', 'ghana': 'accra', 'greece': 'athens',
    'grenada': "saint george's", 'guatemala': 'guatemala city', 'guinea': 'conakry',
    'guinea-bissau': 'bissau', 'equatorial guinea': 'malabo', 'haiti': 'port-au-prince',
    'honduras': 'tegucigalpa', 'hungary': 'budapest', 'india': 'new delhi',
    'indonesia': 'jakarta', 'iran': 'tehran', 'iraq': 'baghdad',
    'ireland': 'dublin', 'israel': 'jerusalem', 'italy': 'rome',
    'jamaica': 'kingston', 'japan': 'tokyo', 'jordan': 'amman',
    'kazakhstan': 'astana', 'kenya': 'nairobi', 'kiribati': 'tarawa',
    'kosovo': 'pristina', 'kuwait': 'kuwait city', 'kyrgyzstan': 'bishkek',
    'laos': 'vientiane', 'latvia': 'riga', 'lebanon': 'beirut',
    'lesotho': 'maseru', 'liberia': 'monrovia', 'libya': 'tripoli',
    'liechtenstein': 'vaduz', 'lithuania': 'vilnius', 'luxembourg': 'luxembourg',
    'madagascar': 'antananarivo', 'malaysia': 'kuala lumpur', 'malawi': 'lilongwe',
    'maldives': 'male', 'mali': 'bamako', 'malta': 'valletta',
    'morocco': 'rabat', 'mauritius': 'port louis', 'mauritania': 'nouakchott',
    'mexico': 'mexico city', 'micronesia': 'palikir', 'moldova': 'chisinau',
    'monaco': 'monaco', 'mongolia': 'ulaanbaatar', 'montenegro': 'podgorica',
    'mozambique': 'maputo', 'myanmar': 'naypyidaw', 'namibia': 'windhoek',
    'nauru': 'yaren', 'nepal': 'kathmandu', 'nicaragua': 'managua',
    'niger': 'niamey', 'nigeria': 'abuja', 'norway': 'oslo',
    'new zealand': 'wellington', 'oman': 'muscat', 'netherlands': 'amsterdam',
    'pakistan': 'islamabad', 'palau': 'melekeok', 'panama': 'panama city',
    'papua new guinea': 'port moresby', 'paraguay': 'asuncion', 'peru': 'lima',
    'poland': 'warsaw', 'portugal': 'lisbon',
    'romania': 'bucharest', 'russia': 'moscow', 'rwanda': 'kigali',
    'samoa': 'apia', 'san marino': 'san marino', 'sao tome and principe': 'sao tome',
    'serbia': 'belgrade', 'seychelles': 'victoria', 'sierra leone': 'freetown',
    'singapore': 'singapore', 'syria': 'damascus', 'south africa': 'pretoria',
    'sudan': 'khartoum', 'south sudan': 'juba', 'sweden': 'stockholm',
    'switzerland': 'bern', 'thailand': 'bangkok', 'tanzania': 'dodoma',
    'togo': 'lome', 'tonga': "nuku'alofa", 'tunisia': 'tunis',
    'turkey': 'ankara', 'turkmenistan': 'ashgabat', 'tuvalu': 'funafuti',
    'ukraine': 'kyiv', 'uganda': 'kampala', 'united kingdom': 'london',
    'united states': 'washington, d.c.', 'uruguay': 'montevideo',
    'uzbekistan': 'tashkent', 'vanuatu': 'port vila', 'vatican city': 'vatican city',
    'venezuela': 'caracas', 'vietnam': 'hanoi', 'yemen': "sana'a",
    'zambia': 'lusaka', 'zimbabwe': 'harare'
}


def read_int(prompt: str, minimum: int = 1) -> int:
    # Read an integer from the user with validation
    while True:
        raw = input(prompt).strip()
        try:
            value = int(raw)
            if value < minimum:
                print(f"Please enter a number >= {minimum}.")
                continue
            return value
        except ValueError:
            print("Please enter a valid integer.")


def get_output_folder() -> Path:
    # Ask for an output folder; default to ./output_quizzes if blank
    raw = input("Enter output folder path (leave empty for ./output_quizzes): ").strip()
    if not raw:
        return Path.cwd() / "output_quizzes"
    return Path(raw)


def get_output_format() -> str:
    # Ask the user for the desired output format
    while True:
        fmt = input("Choose output format (txt / pdf / docx): ").strip().lower()
        if fmt in ("txt", "pdf", "docx"):
            return fmt
        print("Invalid format. Please choose: txt, pdf, or docx.")


def wrap_line(text: str, max_chars: int) -> list[str]:
    # Simple word-wrap for PDF (monospace)
    if len(text) <= max_chars:
        return [text]

    words = text.split(" ")
    lines: list[str] = []
    current = ""

    for w in words:
        if not current:
            current = w
            continue

        if len(current) + 1 + len(w) <= max_chars:
            current += " " + w
        else:
            lines.append(current)
            current = w

    if current:
        lines.append(current)

    return lines


def build_quiz_text(selected_countries: list[str]) -> tuple[str, str]:
    # Build quiz text and answer key text in memory (no title here to avoid duplicates)
    quiz_lines: list[str] = []
    answer_lines: list[str] = []

    quiz_lines.append("Name: ________________________________")
    quiz_lines.append("Date: _________________________________")
    quiz_lines.append("Period: _______________________________")
    quiz_lines.append("")

    # For each question, generate 1 correct + 3 wrong answers (unique)
    for q_index, country in enumerate(selected_countries, start=1):
        correct = capitals[country]

        # Use a unique pool to avoid duplicate options
        wrong_pool = list(set(capitals.values()) - {correct})
        wrong = random.sample(wrong_pool, 3)

        options = wrong + [correct]
        random.shuffle(options)

        quiz_lines.append(f"{q_index}. What is the capital of {country}?")
        for i, option in enumerate(options):
            quiz_lines.append(f"    {'ABCD'[i]}. {option}")
        quiz_lines.append("")

        correct_letter = "ABCD"[options.index(correct)]
        answer_lines.append(f"{q_index}. {correct_letter}")

    return "\n".join(quiz_lines) + "\n", "\n".join(answer_lines) + "\n"


def save_txt(path: Path, title: str, content: str) -> None:
    # Save plain text with title at the top (so TXT also has a title)
    with open(path, "w", encoding="utf-8") as f:
        f.write(title + "\n\n")
        f.write(content)


def save_docx(path: Path, title: str, content: str) -> None:
    # Save DOCX with a clean heading + monospace body
    doc = Document()
    doc.add_heading(title, level=1)

    # Monospace for consistent alignment
    style = doc.styles["Normal"]
    style.font.name = "Courier New"
    style.font.size = Pt(10)

    for line in content.splitlines():
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(0)

    doc.save(path)


def save_pdf(path: Path, title: str, content: str) -> None:
    # Save a readable PDF using reportlab canvas (monospace + pagination)
    c = canvas.Canvas(str(path))
    page_width, page_height = c._pagesize

    left_margin = 50
    top_margin = 50
    bottom_margin = 50

    title_font = "Helvetica-Bold"
    body_font = "Courier"
    title_size = 16
    body_size = 10

    line_height = body_size + 4

    # Estimate max characters per line in monospace for this page width
    usable_width = page_width - (left_margin * 2)
    max_chars = max(40, int(usable_width / (body_size * 0.6)))

    def new_page():
        c.showPage()
        c.setFont(body_font, body_size)

    # Title
    y = page_height - top_margin
    c.setFont(title_font, title_size)
    c.drawString(left_margin, y, title)

    # Body
    y -= (title_size + 14)
    c.setFont(body_font, body_size)

    for raw_line in content.splitlines():
        # Wrap long lines for PDF
        lines = wrap_line(raw_line, max_chars) if raw_line.strip() else [""]

        for line in lines:
            if y <= bottom_margin:
                new_page()
                y = page_height - top_margin
            c.drawString(left_margin, y, line)
            y -= line_height

    c.save()


def main() -> None:
    # Requirements (install once):
    # python -m pip install python-docx reportlab

    output_folder = get_output_folder()
    output_format = get_output_format()

    num_quizzes = read_int("Enter the number of quizzes to generate: ", minimum=1)
    num_questions = read_int("Enter the number of questions per quiz: ", minimum=1)

    if num_questions > len(capitals):
        raise ValueError(
            f"Cannot generate {num_questions} questions. "
            f"Only {len(capitals)} countries are available."
        )

    output_folder.mkdir(parents=True, exist_ok=True)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    all_countries = list(capitals.keys())

    print(f"\nSaving files to: {output_folder}")
    print(f"Format: {output_format}\n")

    for quiz_num in range(1, num_quizzes + 1):
        # Shuffle a fresh copy so each quiz is independent
        countries_copy = all_countries[:]
        random.shuffle(countries_copy)
        selected = countries_copy[:num_questions]

        quiz_text, answers_text = build_quiz_text(selected)

        base_name = f"capital_quiz_{stamp}_{quiz_num}"
        quiz_title = f"World Capitals Quiz (Form {quiz_num})"
        answers_title = f"Answer Key - World Capitals Quiz (Form {quiz_num})"

        if output_format == "txt":
            quiz_path = output_folder / f"{base_name}.txt"
            answers_path = output_folder / f"{base_name}_answers.txt"
            save_txt(quiz_path, quiz_title, quiz_text)
            save_txt(answers_path, answers_title, answers_text)

        elif output_format == "docx":
            quiz_path = output_folder / f"{base_name}.docx"
            answers_path = output_folder / f"{base_name}_answers.docx"
            save_docx(quiz_path, quiz_title, quiz_text)
            save_docx(answers_path, answers_title, answers_text)

        else:  # pdf
            quiz_path = output_folder / f"{base_name}.pdf"
            answers_path = output_folder / f"{base_name}_answers.pdf"
            save_pdf(quiz_path, quiz_title, quiz_text)
            save_pdf(answers_path, answers_title, answers_text)

        print(f"Created: {quiz_path.name}")
        print(f"Created: {answers_path.name}\n")

    print("✔️ Quizzes generated successfully.")


if __name__ == "__main__":
    main()
